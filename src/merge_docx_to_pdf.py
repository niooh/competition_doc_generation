"""
合并 docx 试题并生成带封面、目录、水印的 PDF。
年份由项目根目录 config.py 的 YEAR 变量配置。
"""

import os
import time
import tempfile
from pathlib import Path
from typing import List, Tuple

import pythoncom
import win32com.client as win32
from pypdf import PdfReader, PdfWriter
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import mm
from reportlab.pdfgen import canvas
from reportlab.lib.colors import gray
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from docx import Document

# 读取年份
from src.config import YEAR

# 候选中文字体路径，按优先级排列
_CANDIDATE_FONT_PATHS: List[Tuple[str, str]] = [
    ("Microsoft YaHei", "C:/Windows/Fonts/msyh.ttc"),
    ("Microsoft YaHei Bold", "C:/Windows/Fonts/msyhbd.ttc"),
    ("SimSun", "C:/Windows/Fonts/simsun.ttc"),
    ("SimHei", "C:/Windows/Fonts/simhei.ttf"),
    ("KaiTi", "C:/Windows/Fonts/simkai.ttf"),
]


def register_cjk_font(font_name: str = "CJKFont") -> str:
    """注册第一个可用的中文字体，返回字体名；失败则抛出 FileNotFoundError。"""
    for family, path in _CANDIDATE_FONT_PATHS:
        if os.path.isfile(path):
            pdfmetrics.registerFont(TTFont(font_name, path))
            print(f"Registered CJK font: {family}")
            return font_name
    raise FileNotFoundError(
        "No CJK font found. Install SimSun, Microsoft YaHei or SimHei, "
        "or update _CANDIDATE_FONT_PATHS."
    )


CJK_FONT = register_cjk_font()


def convert_one(docx_path: str, pdf_path: str) -> bool:
    """单个 docx 转 PDF，失败自动重试 2 次，返回成功与否。"""
    max_retries = 2
    for attempt in range(max_retries):
        word = None
        doc = None
        try:
            pythoncom.CoInitialize()
            word = win32.Dispatch("Word.Application")
            word.Visible = False
            word.DisplayAlerts = 0
            doc = word.Documents.Open(docx_path, ReadOnly=True, Visible=False)
            doc.SaveAs(pdf_path, FileFormat=17)
            return True
        except Exception as e:
            print(f"  ✗ Attempt {attempt + 1}/{max_retries} failed: {e}")
            if attempt == max_retries - 1:
                return False
            time.sleep(1)
        finally:
            try:
                if doc:
                    doc.Close(False)
            except Exception:
                pass
            try:
                if word:
                    word.Quit()
            except Exception:
                pass
            pythoncom.CoUninitialize()
    return False


def create_cover_pdf(year: int, output_path: str) -> None:
    """生成封面页，标题居中。"""
    c = canvas.Canvas(output_path, pagesize=A4)
    width, height = A4
    title = f"{year}年暑假化学吧吧赛试题（大学组）"
    c.setFont(CJK_FONT, 28)
    c.drawCentredString(width / 2, height / 2, title)
    c.save()


def create_toc_pdf(
    entries: List[Tuple[str, int]], body_offset: int, output_path: str
) -> None:
    """生成目录 PDF，显示文档标题与最终页码。"""
    c = canvas.Canvas(output_path, pagesize=A4)
    width, height = A4
    left_margin = 40 * mm
    right_margin = width - 40 * mm
    y = height - 30 * mm
    line_height = 7 * mm

    c.setFont(CJK_FONT, 12)
    for title, body_page in entries:
        final_page = body_page + body_offset - 1
        if y < 30 * mm:
            c.showPage()
            c.setFont(CJK_FONT, 12)
            y = height - 30 * mm
        c.drawString(left_margin, y, title[:80])
        c.drawRightString(right_margin, y, str(final_page))
        y -= line_height
    c.save()


def create_watermark_page(
    year: int, page_width: float, page_height: float, output_path: str
) -> None:
    """生成灰色倾斜水印单页。"""
    c = canvas.Canvas(output_path, pagesize=(page_width, page_height))
    c.setFillColor(gray, alpha=0.15)
    c.setFont(CJK_FONT, 72)
    c.translate(page_width / 2, page_height / 2)
    c.rotate(45)
    c.drawCentredString(0, 0, f"化吧吧赛 {year}")
    c.save()


def add_watermark_to_pdf(input_pdf: str, output_pdf: str, year: int) -> None:
    """将水印叠加到 PDF 的所有页面。"""
    reader = PdfReader(input_pdf)
    writer = PdfWriter()

    page0 = reader.pages[0]
    w = float(page0.mediabox.width)
    h = float(page0.mediabox.height)

    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
        watermark_path = tmp.name
    create_watermark_page(year, w, h, watermark_path)
    watermark_page = PdfReader(watermark_path).pages[0]

    for page in reader.pages:
        page.merge_page(watermark_page, over=True)  # type: ignore[arg-type]
        writer.add_page(page)

    with open(output_pdf, "wb") as f:
        writer.write(f)
    os.unlink(watermark_path)


def merge_docx_to_pdf(input_dir: Path, output_pdf: Path) -> None:
    """主流程：合并 docx，附加封面/目录/水印，输出最终 PDF。"""
    input_dir = input_dir.resolve()
    output_pdf = output_pdf.resolve()
    temp_dir = output_pdf.parent / "temp_pdf"
    temp_dir.mkdir(exist_ok=True)

    docx_files = sorted(input_dir.glob("*.docx"))
    if not docx_files:
        print("No .docx files found")
        return

    pdf_files: List[Path] = []
    print(f"Converting {len(docx_files)} files...")
    start = time.perf_counter()

    for docx_path in docx_files:
        pdf_path = temp_dir / (docx_path.stem + ".pdf")
        print(f"→ {docx_path.name}")
        if convert_one(str(docx_path), str(pdf_path)):
            pdf_files.append(pdf_path)
            print("  ✓ Done")
        else:
            print("  ✗ Failed")

    if not pdf_files:
        print("No files converted")
        return

    elapsed = time.perf_counter() - start
    print(f"Conversion finished in {elapsed:.1f}s")

    # 合并试题 PDF，记录各文档标题及起始页码
    body_writer = PdfWriter()
    doc_info: List[Tuple[str, int]] = []
    for fpath, docx_path in zip(pdf_files, docx_files):
        reader = PdfReader(str(fpath))
        start_page = len(body_writer.pages) + 1
        try:
            doc = Document(str(docx_path))
            title = doc.paragraphs[0].text.strip() if doc.paragraphs else docx_path.stem
        except Exception:
            title = docx_path.stem
        doc_info.append((title, start_page))
        for page in reader.pages:
            body_writer.add_page(page)

    body_pdf = temp_dir / "body.pdf"
    with open(body_pdf, "wb") as f:
        body_writer.write(f)

    cover_pdf = temp_dir / "cover.pdf"
    create_cover_pdf(YEAR, str(cover_pdf))

    # 目录：先估算页数，获取实际页数后若不符则重生成
    assumed_toc_pages = 1
    body_offset_guess = 1 + assumed_toc_pages
    toc_temp = temp_dir / "toc_temp.pdf"
    create_toc_pdf(doc_info, body_offset_guess, str(toc_temp))
    toc_reader = PdfReader(str(toc_temp))
    real_toc_pages = len(toc_reader.pages)
    if real_toc_pages != assumed_toc_pages:
        body_offset = 1 + real_toc_pages
        toc_final = temp_dir / "toc.pdf"
        create_toc_pdf(doc_info, body_offset, str(toc_final))
        os.unlink(toc_temp)
    else:
        toc_final = toc_temp

    # 合并封面 + 目录 + 试题
    merged_no_wm = temp_dir / "merged_no_watermark.pdf"
    merger = PdfWriter()
    for pdf in [cover_pdf, toc_final, body_pdf]:
        reader = PdfReader(str(pdf))
        for page in reader.pages:
            merger.add_page(page)
    with open(merged_no_wm, "wb") as f:
        merger.write(f)

    add_watermark_to_pdf(str(merged_no_wm), str(output_pdf), YEAR)

    # 清理临时文件
    for p in pdf_files + [body_pdf, cover_pdf, toc_final, merged_no_wm]:
        try:
            os.unlink(p)
        except Exception:
            pass
    if toc_temp.exists():
        try:
            os.unlink(toc_temp)
        except Exception:
            pass
    try:
        temp_dir.rmdir()
    except Exception:
        pass

    print(f"✓ Output saved to {output_pdf}")


def main() -> None:
    base = Path(__file__).resolve().parent.parent
    dist_dir = base / "dist"
    dist_dir.mkdir(exist_ok=True)
    merge_docx_to_pdf(base / "doc", dist_dir / "merged.pdf")


if __name__ == "__main__":
    main()
