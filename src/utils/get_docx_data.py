from pathlib import Path
from docx import Document
from docx.document import Document as DocumentType
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
import os

from src.config import DOC_DIR


def extract_end_score(line: str) -> int | None:
    """提取行尾分值，如 '（25分）'，返回整数或 None。"""
    s = line.rstrip()
    if not s:
        return None

    last_left = s.rfind('（')
    if last_left == -1:
        return None
    right = s.find('）', last_left)
    if right == -1:
        return None

    content = s[last_left + 1: right]
    fen_pos = content.find('分')
    if fen_pos <= 0:
        return None

    num_str = ''
    i = fen_pos - 1
    while i >= 0 and content[i].isdigit():
        num_str = content[i] + num_str
        i -= 1

    if not num_str:
        return None

    if right != len(s) - 1:
        print(f"! 分数括号不在行尾 | 行内容：{s}")

    return int(num_str)


def get_paragraph_level(para: Paragraph) -> int:
    """获取 Word 大纲级别（无则为0）。"""
    num_pr = para._p.xpath('.//w:numPr')
    if not num_pr:
        return 0
    ilvl = num_pr[0].find(qn('w:ilvl'))
    if ilvl is None:
        return 0
    return int(ilvl.get(qn('w:val'))) + 1


def get_doc_title(doc: DocumentType) -> str:
    """取文档第一个非空段落作为标题。"""
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            return t
    return "undefined"


def filter_parent_items(items: list[tuple[str, int]]) -> list[tuple[str, int]]:
    """过滤掉单层编号（如'1'、'2'），只保留子级（如'1-1'）。"""
    return [item for item in items if '-' in item[0]]


def parse_doc(doc_path: str) -> tuple[str, list[tuple[str, int]]]:
    """解析单个文档，返回 (标题, [(编号, 分数)])。"""
    doc = Document(doc_path)
    title = get_doc_title(doc)
    result = []
    level_counters = [0]
    prev_level = 0

    for p in doc.paragraphs:
        txt = p.text.strip()
        if not txt:
            continue
        curr_level = get_paragraph_level(p)
        score = extract_end_score(txt)
        if curr_level <= 0:
            continue

        if curr_level > prev_level:
            while len(level_counters) < curr_level:
                level_counters.append(0)
        elif curr_level < prev_level:
            level_counters = level_counters[:curr_level]

        level_counters[curr_level - 1] += 1
        num_str = "-".join(map(str, level_counters[:curr_level]))
        if score is not None:
            result.append((num_str, score))
        prev_level = curr_level

    return title, filter_parent_items(result)


def docx_data_to_ls() -> dict[str, list[tuple[str, int]]]:
    """批量读取 DOC_DIR 下所有 docx，返回 {标题: [(编号, 分数)]}。"""
    folder = DOC_DIR
    res = {}
    if not folder.exists():
        folder.mkdir(parents=True)
        print(f"⚠️ 创建了 doc/，请将 .docx 放入后重新运行")
        return res

    for fname in os.listdir(folder):
        if fname.lower().endswith(".docx"):
            fp = folder / fname
            try:
                title, filtered = parse_doc(str(fp))
                res[title] = filtered
            except Exception as e:
                print(f"✗ {fname}：{e}")
    return res


if __name__ == "__main__":
    data = docx_data_to_ls()
    for t, content in data.items():
        print(f"{t}: {content}")
