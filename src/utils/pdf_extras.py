import io
from pathlib import Path

from reportlab.lib.pagesizes import A4
from reportlab.lib.units import mm
from reportlab.pdfgen import canvas
from pypdf import PdfReader, PdfWriter
from docx import Document
from PIL import Image

from src.config import YEAR, GROUP, TEMP_DIR, TOC_TOP_MARGIN

from src.utils.fonts import get_cjk_font

def add_page_numbers(input_pdf: str, output_pdf: str) -> None:
    """给 PDF 每页底部居中添加从 1 开始的页码"""
    reader = PdfReader(input_pdf)
    writer = PdfWriter()
    font_name = get_cjk_font()
    for i, page in enumerate(reader.pages):
        w = float(page.mediabox.width)
        h = float(page.mediabox.height)
        packet = io.BytesIO()
        c = canvas.Canvas(packet, pagesize=(w, h))
        c.setFont(font_name, 10)
        c.drawCentredString(w / 2, 20, str(i + 1))
        c.save()
        packet.seek(0)
        overlay = PdfReader(packet).pages[0]
        page.merge_page(overlay, over=True)
        writer.add_page(page)
    with open(output_pdf, "wb") as f:
        writer.write(f)


def create_cover(year: int = YEAR, group: str = GROUP, output_path: str | Path | None = None) -> str:
    """生成封面页，返回临时文件路径（若未指定则使用默认临时路径）。"""
    if output_path is None:
        TEMP_DIR.mkdir(parents=True, exist_ok=True)
        output_path = TEMP_DIR / "cover.pdf"
    else:
        output_path = Path(output_path)

    c = canvas.Canvas(str(output_path), pagesize=A4)
    width, height = A4
    line1 = f"{year}年暑假化学吧吧赛试题"
    line2 = f"（{group}）"
    c.setFont(get_cjk_font(), 28)
    # 第一行放在垂直中点偏上的位置，第二行放在偏下的位置
    c.drawCentredString(width / 2, height / 2 + 24, line1)
    c.drawCentredString(width / 2, height / 2 - 16, line2)

    c.save()
    return str(output_path)


def create_toc(entries: list[tuple[str, int]],
               image_path: str | Path | None = None,
               output_path: str | Path | None = None) -> str:
    """
    生成目录PDF，返回文件路径。
    entries: [(文档标题, 在正文中的起始页码), ...]
    image_path: 可选，在目录下方插入的图片（如周期表）路径
    """
    if output_path is None:
        TEMP_DIR.mkdir(parents=True, exist_ok=True)
        output_path = TEMP_DIR / "toc.pdf"
    else:
        output_path = Path(output_path)

    c = canvas.Canvas(str(output_path), pagesize=A4)
    width, height = A4

    # 页面布局参数（可根据需要调整）
    left_margin = 30 * mm
    right_margin = width - 30 * mm
    toc_font_size = 16
    line_height = 13 * mm
    # 底部最小保留空间
    bottom_margin = 8 * mm

    font_name = get_cjk_font()
    c.setFont(font_name, toc_font_size)

    y = height - TOC_TOP_MARGIN * mm

    # 绘制目录条目
    for title, body_page in entries:
        # 如果当前页剩余空间不足以再画一行，换页
        if y < bottom_margin + line_height:
            c.showPage()
            c.setFont(font_name, toc_font_size)
            y = height - bottom_margin

        title_text = title[:80]
        c.drawString(left_margin, y, title_text)
        c.drawRightString(right_margin, y, str(body_page))
        y -= line_height

    # 在目录下方绘制图片（如周期表）
    if image_path is not None and Path(image_path).is_file():
        img_display_width = width * 0.8  # 图片占页面宽度的比例
        with Image.open(image_path) as img:
            img_w, img_h = img.size
        ratio = img_display_width / img_w
        img_display_height = img_h * ratio

        gap_above_img = 10 * mm
        # 图片所需总高度：图片本身 + 上方间距 + 底部留白
        required_height = img_display_height + gap_above_img + bottom_margin

        # 如果剩余空间放不下图片，新开一页
        if y - required_height < bottom_margin:
            c.showPage()
            y = height - bottom_margin

        # 居中放置图片
        img_x = (width - img_display_width) / 2
        img_y = y - gap_above_img - img_display_height
        c.drawImage(str(image_path), img_x, img_y,
                    width=img_display_width, height=img_display_height)

    c.save()
    return str(output_path)
    
def merge_pdfs(pdf_paths: list[str | Path], output_path: str | Path) -> None:
    """按顺序合并多个PDF并写入 output_path。"""
    writer = PdfWriter()
    for p in pdf_paths:
        reader = PdfReader(str(p))
        for page in reader.pages:
            writer.add_page(page)
    with open(output_path, "wb") as f:
        writer.write(f)


def get_doc_title(docx_path: str | Path) -> str:
    """从 docx 第一段提取标题，失败则返回文件名。"""
    try:
        doc = Document(str(docx_path))
        if doc.paragraphs:
            return doc.paragraphs[0].text.strip()
    except Exception:
        pass
    return Path(docx_path).stem
