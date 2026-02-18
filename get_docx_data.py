from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
import os

# 提取行尾分值（无分返回None）
def extract_end_score(line: str) -> int | None:
    s = line.rstrip()
    if not s:
        return None

    last_left_bracket = s.rfind('（')  # 从后往前找最后一个（
    if last_left_bracket == -1:
        return None
    right_bracket = s.find('）', last_left_bracket)  # 从这里往后找第一个 ）
    if right_bracket == -1:
        return None

    # 截取括号内部内容：（ content ）
    content = s[last_left_bracket + 1 : right_bracket]

    fen_pos = content.find('分')
    if fen_pos <= 0:
        return None

    # 分前连续数字
    num_str = ''
    i = fen_pos - 1
    while i >= 0 and content[i].isdigit():
        num_str = content[i] + num_str
        i -= 1

    if not num_str:
        return None

    # 判断：这个括号是否在行尾
    if right_bracket != len(s) - 1:
        print(f"! 分数括号不在行尾 | 行内容：{line.strip()}")

    return int(num_str)


# 获取Word真实层级（无大纲层级返回0）
def get_paragraph_level(para: Paragraph) -> int:
    num_pr = para._p.xpath('.//w:numPr')
    if not num_pr:
        return 0
    ilvl = num_pr[0].find(qn('w:ilvl'))
    if ilvl is None:
        return 0
    return int(ilvl.get(qn('w:val'))) + 1

# 取文档标题
def get_doc_title(doc: Document) -> str:
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            return t
    return "undefined"

# 对原[(编号, 分数), ...]列表，过滤掉单层编号的父级（如1/2/3），只保留子级（如1-1/3-2-1）
def filter_parent_items(origin_list: list[tuple[str, int]]) -> list[tuple[str, int]]:
    return [item for item in origin_list if '-' in item[0]]

# 解析文档：匹配Word实际编号，兼容无分结构级
def parse_doc_correct(doc_path: str) -> tuple[str, list[tuple[str, int]]]:
    doc = Document(doc_path)
    title = get_doc_title(doc)
    result = []
    level_counters = [0]  # 层级计数器，索引=层级-1
    prev_level = 0  # 上一个有效大纲层级

    for p in doc.paragraphs:
        txt = p.text.strip()
        if not txt:
            continue
        curr_level = get_paragraph_level(p)
        score = extract_end_score(txt)
        if curr_level <= 0:
            continue

        # 层级处理逻辑
        if curr_level > prev_level:
            while len(level_counters) < curr_level:
                level_counters.append(0)
        elif curr_level < prev_level:
            level_counters = level_counters[:curr_level]

        # 生成编号并收集结果（有分才加）
        level_counters[curr_level - 1] += 1
        num_str = "-".join(map(str, level_counters[:curr_level]))
        if score is not None:
            result.append((num_str, score))
        
        prev_level = curr_level

    # 解析后直接过滤父级，返回子级列表
    filtered_result = filter_parent_items(result)
    return title, filtered_result

# 批量读取doc文件夹下所有docx的数据，得到列表
def docx_data_to_ls() -> dict[str, list[tuple[str, int]]]:
    folder = "./doc/"
    res = {}
    if not os.path.exists(folder):
        os.makedirs(folder)
        print(f"⚠️  创建了doc文件夹，请将Word文件放入后重新运行")
        return res

    for fname in os.listdir(folder):
        if fname.lower().endswith(".docx"):
            fp = os.path.join(folder, fname)
            try:
                title, filtered_data = parse_doc_correct(fp)
                res[title] = filtered_data
                #print(f"✓ {fname}")
                #for item in filtered_data: print(f"   {item}")
            except Exception as e:
                print(f"❌ {fname}：{str(e)}\n")
    return res

if __name__ == "__main__":
    final_data = docx_data_to_ls()
    for doc_title, content in final_data.items():
        print(f"{doc_title}: {content}")
